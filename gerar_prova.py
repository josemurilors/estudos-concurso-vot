#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Gera prova estilo Vunesp - 100 questões: Cisco IOS, Mikrotik, Linux CentOS."""

import random
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

# =======================================================================
# BANCO DE QUESTÕES
# Cada item: (enunciado, [A,B,C,D,E], índice_correto(0-4), comentário, tema)
# =======================================================================

Q = []

# ----------------- CISCO IOS (34 questões) -----------------
Q.append(("Em um roteador Cisco IOS, o administrador precisa configurar uma rota estática padrão apontando para o próximo salto 200.10.1.1. O comando correto, em modo de configuração global, é:",
["ip route default 200.10.1.1","ip route 0.0.0.0 0.0.0.0 200.10.1.1","ip default-gateway 200.10.1.1","route add default 200.10.1.1","ip static-route 0.0.0.0/0 200.10.1.1"],1,
"A rota estática padrão no Cisco IOS é definida com 'ip route 0.0.0.0 0.0.0.0 <next-hop>'. O comando 'ip default-gateway' só é válido quando o roteamento está desativado (no ip routing), típico em switches L2.","cisco"))

Q.append(("Considere a seguinte configuração em um switch Cisco: \ninterface FastEthernet0/1\n switchport mode access\n switchport access vlan 10\n switchport port-security\n switchport port-security maximum 2\n switchport port-security violation restrict\nAssinale a alternativa correta sobre o comportamento configurado:",
["Ao exceder 2 endereços MAC, a porta é desligada (err-disabled).","Ao exceder 2 endereços MAC, os quadros excedentes são descartados silenciosamente, sem incremento de contador.","Ao exceder 2 endereços MAC, os quadros excedentes são descartados, o contador de violação é incrementado e uma notificação SNMP pode ser gerada, mas a porta permanece ativa.","A porta aceita qualquer número de MACs desde que estejam na VLAN 10.","O modo 'restrict' equivale ao modo 'shutdown', apenas com mensagem de log diferente."],2,
"No modo 'restrict' os quadros violadores são descartados, os contadores são incrementados e pode ser gerado trap SNMP/syslog, sem colocar a porta em err-disabled (o que ocorreria no modo 'shutdown').","cisco"))

Q.append(("No protocolo OSPF, a métrica padrão utilizada pelo Cisco IOS é calculada com base em:",
["Número de saltos (hop count).","Largura de banda de referência dividida pela largura de banda da interface.","Atraso acumulado do caminho.","Largura de banda somada ao atraso e carga.","MTU da interface."],1,
"O custo OSPF no Cisco IOS é, por padrão, 10^8 / bandwidth(bps). A métrica combinada com BW+Delay+Load+Reliability é a do EIGRP.","cisco"))

Q.append(("No EIGRP, o termo 'Feasible Successor' designa:",
["A rota ativa atualmente utilizada para o destino.","Uma rota reserva que satisfaz a condição de viabilidade (Reported Distance < Feasible Distance do successor).","Qualquer rota presente na tabela de topologia.","A rota cujo custo é igual ao do successor.","Uma rota aprendida via redistribuição."],1,
"O Feasible Successor é a rota de backup que satisfaz a Feasibility Condition: a Reported Distance (RD) do vizinho é menor que a Feasible Distance (FD) atual.","cisco"))

Q.append(("Qual comando permite visualizar as tabelas de vizinhos OSPF em um roteador Cisco?",
["show ip ospf database","show ip ospf neighbor","show ospf adjacency","show ip protocols ospf","show neighbor ospf"],1,
"'show ip ospf neighbor' lista os vizinhos OSPF e o estado da adjacência (Full, 2-Way, etc.).","cisco"))

Q.append(("Em um tronco 802.1Q no Cisco IOS, a VLAN nativa:",
["É obrigatoriamente a VLAN 1 e não pode ser alterada.","Tem seu tráfego marcado (tagged) com o VID 1 no tronco.","Tem seu tráfego trafegado sem marcação (untagged) no tronco.","Corresponde à VLAN de gerenciamento por padrão.","É usada somente para o protocolo VTP."],2,
"No 802.1Q o tráfego da VLAN nativa trafega sem tag no tronco. Por padrão é a VLAN 1, mas pode ser alterada com 'switchport trunk native vlan <id>'.","cisco"))

Q.append(("A seguinte ACL estendida foi aplicada em R1:\naccess-list 101 deny tcp 192.168.10.0 0.0.0.255 any eq 23\naccess-list 101 permit ip any any\nAssinale a alternativa correta:",
["Bloqueia todo o tráfego da rede 192.168.10.0/24.","Bloqueia apenas o tráfego Telnet originado em 192.168.10.0/24 e libera o restante.","Bloqueia o tráfego Telnet destinado a 192.168.10.0/24.","Libera apenas Telnet para a rede 192.168.10.0/24.","É inválida por misturar ACL padrão e estendida."],1,
"A regra nega TCP com porta destino 23 (Telnet) vinda da rede 192.168.10.0/24; em seguida permite o restante (permit ip any any).","cisco"))

Q.append(("No STP (Spanning Tree Protocol IEEE 802.1D), a eleição da Root Bridge considera:",
["O maior endereço MAC da switch.","A menor Bridge ID (prioridade + MAC).","A maior prioridade configurada.","A interface com maior custo.","A quantidade de VLANs configuradas."],1,
"A Bridge ID mais baixa (Priority + MAC) vence a eleição de Root Bridge. Menor prioridade = mais preferencial.","cisco"))

Q.append(("Em um ambiente com VTP (VLAN Trunking Protocol) versão 2, qual modo permite criar, modificar e excluir VLANs, mas NÃO propaga essas alterações para outros switches?",
["Server","Client","Transparent","Off","Forwarding"],2,
"No modo Transparent as mudanças ficam locais; entretanto, o switch encaminha (forward) as mensagens VTP recebidas de outros. Não cria/edita VLANs o Client; somente o Server propaga.","cisco"))

Q.append(("Para configurar um EtherChannel (LAG) com LACP ativo em duas interfaces de um switch Cisco, utiliza-se:",
["channel-group 1 mode on","channel-group 1 mode desirable","channel-group 1 mode active","channel-group 1 mode auto","channel-group 1 mode passive-auto"],2,
"LACP utiliza os modos 'active' (negocia ativamente) e 'passive'. 'desirable/auto' são do PAgP, protocolo proprietário Cisco. 'on' força sem negociação.","cisco"))

Q.append(("No HSRP (Hot Standby Router Protocol), o roteador com maior valor de prioridade torna-se ativo, desde que:",
["O comando 'standby preempt' esteja configurado nos roteadores; caso contrário, o primeiro roteador ativo permanece, independentemente da prioridade.","A prioridade seja obrigatoriamente 255.","Esteja com interface em modo passivo.","Possua o menor endereço IP.","Seja o único em VLAN nativa."],0,
"Sem 'standby preempt', o router com maior prioridade não toma automaticamente o papel de ativo quando um outro já está ativo.","cisco"))

Q.append(("O comando 'show ip interface brief' em um roteador Cisco exibe, principalmente:",
["Estatísticas detalhadas de cada interface.","Tabela de roteamento resumida.","Lista das interfaces com IP, status administrativo e status de protocolo.","ACLs aplicadas em cada interface.","Lista de VLANs configuradas."],2,
"'show ip interface brief' fornece visão resumida: nome da interface, IP, método, Status e Protocol.","cisco"))

Q.append(("Sobre o BGP (Border Gateway Protocol), é correto afirmar que:",
["É um protocolo de roteamento interno (IGP) baseado em estado de enlace.","Utiliza a porta UDP 179 para estabelecer sessões.","É um protocolo de roteamento entre sistemas autônomos (EGP) que utiliza TCP/179.","Calcula rotas pelo algoritmo de Dijkstra.","Não suporta políticas de roteamento."],2,
"BGP é Path-Vector, entre ASs (EGP), usa TCP porta 179 e é orientado por políticas (atributos como AS_PATH, LOCAL_PREF, MED).","cisco"))

Q.append(("No contexto de NAT no Cisco IOS, qual comando configura o PAT (NAT Overload) utilizando a interface Serial0/0 como endereço global?",
["ip nat inside source list 1 interface Serial0/0 overload","ip nat pool PAT interface Serial0/0","ip nat outside overload Serial0/0","ip nat pat inside Serial0/0","ip nat inside pat Serial0/0 list 1"],0,
"A forma canônica do PAT dinâmico usando a própria interface como endereço global é: 'ip nat inside source list <acl> interface <intf> overload'.","cisco"))

Q.append(("O comando 'copy running-config startup-config' em um dispositivo Cisco IOS tem como efeito:",
["Reiniciar o dispositivo aplicando a configuração atual.","Salvar a configuração em execução (RAM) para a NVRAM.","Copiar a configuração da NVRAM para a RAM.","Exportar a configuração para um servidor TFTP.","Reverter para a configuração de fábrica."],1,
"A running-config fica em RAM e é perdida no reboot. A startup-config é armazenada na NVRAM. O comando grava a RAM na NVRAM.","cisco"))

Q.append(("Considere a configuração:\nrouter ospf 1\n network 10.0.0.0 0.0.0.255 area 0\nIsso significa que:",
["Todas as interfaces com IP na rede 10.0.0.0/24 participarão do OSPF na área 0.","Somente a interface com IP 10.0.0.0 participa.","OSPF ignorará interfaces em 10.0.0.0/24.","A rede 10.0.0.0/24 será anunciada, mas nenhuma interface participará.","A wildcard 0.0.0.255 é inválida para OSPF."],0,
"No OSPF do Cisco IOS usa-se máscara wildcard. 0.0.0.255 equivale a /24; todas as interfaces cujo IP caia nesse intervalo entram no OSPF área 0.","cisco"))

Q.append(("O modo de QoS no Cisco IOS que marca os pacotes no campo DSCP utilizando uma classe definida em class-map é denominado:",
["Policing","Shaping","Marking (MQC - Modular QoS CLI)","Queueing FIFO","Random Early Detection (RED)"],2,
"Dentro de uma policy-map, a ação 'set dscp' aplica marking. MQC é o arcabouço para classificar (class-map) e aplicar políticas (policy-map).","cisco"))

Q.append(("Assinale a alternativa que apresenta corretamente o comando para habilitar o protocolo CDP apenas em uma interface específica:",
["cdp enable (no modo de configuração da interface)","cdp run interface FastEthernet0/1","enable cdp interface","ip cdp enable","cdp advertise-v2"],0,
"No modo de configuração da interface usa-se 'cdp enable'. Globalmente usa-se 'cdp run'.","cisco"))

Q.append(("Para configurar AAA em um roteador Cisco IOS autenticando via servidor RADIUS e mantendo o usuário local como fallback, usa-se:",
["aaa authentication login default group radius local","aaa authentication radius first local second","aaa new-model radius+local","radius-server authenticate local","aaa login group-radius + local"],0,
"A lista 'default' especifica, na ordem, os métodos: primeiro grupo RADIUS, depois base local se o servidor falhar.","cisco"))

Q.append(("O comando 'switchport mode trunk' em um switch Cisco Catalyst força:",
["A porta a negociar modo de tronco via DTP.","A porta a operar em modo acesso caso o vizinho seja acesso.","A porta a operar permanentemente como tronco, enviando DTP para o vizinho.","A porta a operar como tronco sem encapsulamento.","A porta a trocar para modo dynamic desirable."],2,
"'switchport mode trunk' coloca a porta em tronco estático, ainda enviando DTP para o outro lado (a menos que seja usado 'switchport nonegotiate').","cisco"))

Q.append(("No Rapid PVST+ (802.1w), o estado da porta 'Discarding' corresponde, em termos funcionais, aos estados combinados do 802.1D:",
["Listening e Learning","Blocking, Listening e Disabled","Forwarding e Learning","Blocking e Forwarding","Apenas Disabled"],1,
"O estado Discarding do RSTP engloba Disabled, Blocking e Listening do 802.1D, unificando estados de não encaminhamento.","cisco"))

Q.append(("A distância administrativa padrão do OSPF no Cisco IOS é:",
["90","100","110","120","170"],2,
"Valores AD padrão Cisco: EIGRP interno 90, IGRP 100, OSPF 110, IS-IS 115, RIP 120, EIGRP externo 170, iBGP 200, eBGP 20.","cisco"))

Q.append(("No Cisco IOS, o comando 'service password-encryption' tem como característica:",
["Criptografar todas as senhas com SHA-256.","Criptografar as senhas em texto claro do arquivo de configuração usando o algoritmo fraco tipo 7 (Vigenère).","Aplicar hash MD5 a todas as senhas.","Ser equivalente a 'enable secret'.","Desabilitar o login via Telnet."],1,
"O comando aplica o tipo 7 de criptografia (reversível, fraca) apenas para ofuscar senhas em texto claro. Não substitui 'enable secret' que gera hash MD5/SHA.","cisco"))

Q.append(("Sobre o protocolo RIPv2, é correto afirmar:",
["Utiliza broadcast para envio de updates.","Transporta máscaras nas atualizações (classless) e utiliza multicast 224.0.0.9.","Utiliza o algoritmo de Dijkstra.","Possui limite de 255 saltos.","Não utiliza split horizon."],1,
"RIPv2 é classless, envia em multicast 224.0.0.9 via UDP/520, suporta VLSM/autenticação; limite de 15 saltos (16 = infinito).","cisco"))

Q.append(("Qual comando exibe as rotas aprendidas dinamicamente pelo EIGRP no Cisco IOS?",
["show ip route eigrp","show eigrp topology","show eigrp routes","show route eigrp learned","show ip eigrp database"],0,
"'show ip route eigrp' filtra a tabela de roteamento apenas para rotas aprendidas via EIGRP (marcadas com 'D').","cisco"))

Q.append(("Em uma ACL padrão Cisco, o número de identificação deve estar no intervalo:",
["1–99 ou 1300–1999","100–199 ou 2000–2699","200–299","1000–1099","800–999"],0,
"ACLs IP padrão: 1–99 e 1300–1999. ACLs estendidas: 100–199 e 2000–2699.","cisco"))

Q.append(("Para realizar o troubleshooting de conectividade camada 3 entre dois roteadores Cisco, o comando mais indicado para verificar o caminho percorrido pelos pacotes até um destino é:",
["ping","traceroute","debug ip packet","show cdp neighbors","show ip route"],1,
"O traceroute envia pacotes com TTL crescente e exibe cada salto do caminho até o destino.","cisco"))

Q.append(("O tipo de LSA (Link State Advertisement) do OSPF que descreve redes stub e roteadores internos de uma área é o:",
["Tipo 1 - Router LSA","Tipo 2 - Network LSA","Tipo 3 - Summary LSA","Tipo 4 - ASBR Summary","Tipo 5 - External LSA"],0,
"Router LSA (Tipo 1) é gerado por cada roteador dentro da área e descreve as interfaces e seus custos.","cisco"))

Q.append(("Considere a sintaxe: 'ip access-list extended BLOCK_SSH'. Essa configuração:",
["Cria uma ACL padrão nomeada.","Cria uma ACL estendida nomeada 'BLOCK_SSH'.","Aplica imediatamente a ACL em todas as interfaces.","É inválida, pois ACL estendida exige numeração.","Bloqueia SSH automaticamente."],1,
"Permite criar ACL estendida nomeada, com entradas de sequência. Precisa ser aplicada com 'ip access-group BLOCK_SSH in/out' em uma interface.","cisco"))

Q.append(("A tecnologia Inter-VLAN Routing 'Router-on-a-Stick' caracteriza-se por:",
["Utilizar um roteador com várias interfaces físicas, uma por VLAN.","Utilizar subinterfaces em um único link tronco 802.1Q.","Exigir um switch camada 3 com SVIs.","Não permitir uso de 802.1Q.","Criar VLANs dinâmicas via DTP."],1,
"No Router-on-a-Stick configura-se subinterfaces (ex: Gi0/0.10) com 'encapsulation dot1Q <vlan>' em uma interface tronco única.","cisco"))

Q.append(("No Cisco IOS, o comando 'ip helper-address 192.168.1.10' em uma interface:",
["Define um gateway secundário.","Converte broadcasts DHCP/UDP destinados à interface em unicast direcionados ao IP informado.","Cria um túnel GRE.","Ativa NAT dinâmico.","Redireciona SNMP para o endereço informado."],1,
"'ip helper-address' encaminha broadcasts UDP para um servidor (geralmente DHCP) em outra sub-rede. Converte o broadcast em unicast.","cisco"))

Q.append(("Em uma configuração HSRP, o comando 'standby 1 track FastEthernet0/1 30' indica que:",
["A interface FastEthernet0/1 será monitorada e, caso caia, a prioridade será decrementada em 30 unidades.","O roteador aguardará 30 segundos antes de assumir o papel ativo.","A prioridade será fixada em 30.","Serão enviados 30 hellos por segundo.","A VLAN 30 será utilizada para HSRP."],0,
"O 'track' do HSRP reduz a prioridade quando a interface rastreada fica down, permitindo failover automático para o standby.","cisco"))

Q.append(("Sobre o modo EXEC privilegiado no Cisco IOS, é correto afirmar:",
["É acessado com o comando 'configure terminal'.","É identificado pelo prompt terminado em '>' e permite apenas comandos de visualização.","É acessado com o comando 'enable' e permite comandos como reload, debug, copy, etc.","É o modo padrão após a inicialização.","É limitado a consultas em tabelas ARP."],2,
"O modo EXEC Privilegiado é acessado via 'enable' e apresenta o prompt '#'. O modo EXEC Usuário (>) permite apenas consultas básicas.","cisco"))

Q.append(("No Cisco IOS, considere a saída: 'D 10.1.1.0/24 [90/30720] via 10.0.0.2, 00:05:12, GigabitEthernet0/0'. A letra 'D' e os valores entre colchetes indicam, respectivamente:",
["Rota estática / [métrica/custo]","Rota RIP / [hops/delay]","Rota EIGRP / [distância administrativa/métrica composta]","Rota OSPF / [área/cost]","Rota BGP / [peer/as-path]"],2,
"Na tabela de roteamento, 'D' = EIGRP. Entre colchetes está [Distância Administrativa / Métrica]. EIGRP padrão tem AD 90.","cisco"))

# ----------------- MIKROTIK (33 questões) -----------------
Q.append(("No RouterOS (Mikrotik), qual chain padrão da tabela /ip firewall filter é utilizada para tráfego originado ou destinado ao próprio roteador?",
["forward","input","output","srcnat","postrouting"],1,
"A chain 'input' trata pacotes destinados ao roteador; 'output' trata pacotes originados pelo roteador; 'forward' trata tráfego que passa através dele.","mikrotik"))

Q.append(("Em uma regra /ip firewall nat, o action 'masquerade' diferencia-se do 'src-nat' porque:",
["É usado apenas em IPv6.","Altera a porta de origem automaticamente.","Utiliza dinamicamente o endereço IP da interface de saída, sendo indicado para links com IP dinâmico.","É aplicado somente na chain dstnat.","Não realiza tradução de endereço."],2,
"Masquerade é um SNAT que usa o IP atual da interface de saída. Recomendado para conexões com IP dinâmico (PPPoE, DHCP). Em cenários com IP fixo, 'src-nat' é mais eficiente.","mikrotik"))

Q.append(("A ferramenta gráfica oficial do Mikrotik para administração do RouterOS em sistemas Windows é:",
["Putty","Winbox","MacTelnet","SwOS","The Dude"],1,
"Winbox é o utilitário gráfico oficial que se conecta via IP ou MAC (camada 2). The Dude é uma ferramenta de monitoramento. SwOS é um firmware para switches.","mikrotik"))

Q.append(("Qual comando CLI do RouterOS adiciona uma rota estática padrão via gateway 192.168.0.1?",
["/ip route add dst-address=0.0.0.0/0 gateway=192.168.0.1","/ip route default 192.168.0.1","/ip gateway set default=192.168.0.1","/route add default via 192.168.0.1","/routing static add 0.0.0.0 192.168.0.1"],0,
"Sintaxe canônica do Mikrotik: /ip route add dst-address=0.0.0.0/0 gateway=<ip>.","mikrotik"))

Q.append(("No Mikrotik, a tabela 'mangle' do firewall é utilizada principalmente para:",
["Realizar NAT de destino.","Marcar pacotes e conexões (packet-mark, connection-mark, routing-mark) para uso em QoS e roteamento.","Filtrar tráfego malicioso.","Gerenciar listas de usuários Hotspot.","Configurar rotas estáticas."],1,
"O mangle marca pacotes/conexões/rotas para posterior utilização em queues, rotas por política (policy-based routing) e outras funções avançadas.","mikrotik"))

Q.append(("No RouterOS, para criar um servidor PPTP simples, o primeiro passo é:",
["Habilitar o serviço em /interface pptp-server server set enabled=yes","Configurar um bridge com PPTP","Criar uma rota estática para a rede remota","Configurar o IPSec antes do PPTP","Habilitar a interface Winbox"],0,
"O servidor PPTP é ativado em /interface pptp-server server. Depois são criados usuários em /ppp secret e perfil em /ppp profile.","mikrotik"))

Q.append(("O recurso do RouterOS que implementa Simple Queue possibilita:",
["Roteamento avançado via BGP.","Limitação e priorização de banda por IP, interface ou rede, de forma simples, sem necessidade de marcações em mangle.","Criação de VLANs automáticas.","Firewall stateful.","Balanceamento de carga entre links."],1,
"As Simple Queues aplicam limites de upload/download e prioridade por alvo (IP/rede/interface) sem exigir marcação de pacotes.","mikrotik"))

Q.append(("No comando '/ip firewall filter add chain=input protocol=tcp dst-port=22 action=drop in-interface=ether1', a ação esperada é:",
["Bloquear SSH destinado ao próprio roteador apenas quando entra pela ether1.","Bloquear todo tráfego TCP.","Bloquear SSH originado pelo roteador.","Bloquear SSH em trânsito (forward).","Não produz efeito, faltam parâmetros obrigatórios."],0,
"A regra nega pacotes TCP destinados ao roteador (chain=input) na porta 22 (SSH) entrando pela interface ether1.","mikrotik"))

Q.append(("Sobre o User Manager do Mikrotik, é correto afirmar que:",
["É um servidor RADIUS embarcado no RouterOS, que permite autenticação de Hotspot, PPP e Wireless.","É um cliente LDAP.","É usado exclusivamente para HTTP proxy.","Substitui o DNS do sistema.","Gerencia apenas atualizações de firmware."],0,
"User Manager (userman) é um RADIUS integrado para autenticação/accounting de Hotspot, PPPoE/PPTP, Wireless, etc.","mikrotik"))

Q.append(("Para criar uma VLAN de ID 20 sobre a interface ether1 no RouterOS, utiliza-se o comando:",
["/interface vlan add name=vlan20 vlan-id=20 interface=ether1","/ip vlan add id=20 on=ether1","/interface ether1 vlan 20","/vlan add number=20 interface=ether1","/ip switch vlan add 20 ether1"],0,
"A sintaxe correta é /interface vlan add com os parâmetros name, vlan-id e interface.","mikrotik"))

Q.append(("O protocolo de VPN nativamente suportado pelo RouterOS que oferece criptografia forte ponta a ponta, opera nas camadas 3 e utiliza IKE/ESP/AH é o:",
["PPTP","L2TP puro","IPSec","SSTP","EoIP"],2,
"IPSec opera em camada 3, oferecendo confidencialidade (ESP), integridade e autenticação (AH/ESP), com gerenciamento de chaves via IKE/IKEv2.","mikrotik"))

Q.append(("Na CLI do RouterOS, o comando '/export file=backup' gera:",
["Um arquivo binário de backup completo.","Um arquivo texto (.rsc) com todos os comandos de configuração atuais.","Um snapshot da RAM.","Um dump dos pacotes na interface.","Um arquivo de log."],1,
"'/export' produz um arquivo de script texto (.rsc) com a configuração atual em formato de comandos, portável e editável. Para backup binário usa-se '/system backup save'.","mikrotik"))

Q.append(("Sobre o DHCP Server do Mikrotik, ao configurar uma reserva (static lease) é necessário associar:",
["Endereço IP e endereço MAC do cliente.","Endereço IP e hostname.","Apenas o hostname.","Endereço IP e interface física.","Usuário e senha do cliente."],0,
"A reserva (make static) vincula um IP específico ao MAC do cliente, para que ele sempre receba o mesmo endereço via DHCP.","mikrotik"))

Q.append(("No RouterOS, o OSPF é configurado em:",
["/ip ospf","/routing ospf instance e /routing ospf area","/ip routing ospf enable","/interface ospf","/system ospf"],1,
"A partir do RouterOS 6.x/7.x configura-se OSPF em /routing ospf com instance, area e interface-template.","mikrotik"))

Q.append(("Assinale a alternativa que descreve corretamente o recurso 'Hotspot' do Mikrotik:",
["Um proxy reverso para aplicações web.","Um sistema de autenticação transparente (captive portal) que obriga usuários a se autenticarem via página HTML antes de acessar a rede.","Um servidor DNS recursivo.","Uma replicação de rádios wireless.","Um balanceador de carga L4."],1,
"O Hotspot é um captive portal que intercepta conexões HTTP e força o login via página web, integrável com RADIUS/User Manager.","mikrotik"))

Q.append(("A funcionalidade 'Fast Track' do firewall Mikrotik tem como objetivo:",
["Rotear pacotes usando BGP.","Permitir que pacotes de conexões estabelecidas pulem a maior parte do pipeline do firewall/roteamento, aumentando o desempenho (throughput).","Bloquear rapidamente ataques.","Priorizar tráfego VoIP.","Realizar logging rápido em disco."],1,
"FastTrack marca conexões para que pacotes subsequentes sigam um caminho otimizado (bypass de conntrack/firewall), elevando o throughput, porém impossibilitando algumas features em cima desses pacotes.","mikrotik"))

Q.append(("Para configurar uma rede wireless em modo ponto-a-ponto (bridge) no RouterOS, utiliza-se o modo:",
["ap bridge","station bridge (em um lado) e ap bridge (no outro)","alignment-only","station pseudobridge em ambos os lados","nstreme-only"],1,
"Um dos lados opera como 'ap bridge' (ponto de acesso) e o outro como 'station bridge' (cliente capaz de trafegar camada 2). Pseudobridge é para clientes que não suportam station bridge.","mikrotik"))

Q.append(("No RouterOS, o comando '/ip firewall connection print' exibe:",
["As rotas aprendidas.","As conexões ativas rastreadas pelo conntrack, com endereços, portas e estado.","Os logs de firewall.","As interfaces físicas.","Os usuários conectados via PPP."],1,
"Mostra a tabela de estados do conntrack: src/dst, portas, protocolo e estado (established, new, etc.).","mikrotik"))

Q.append(("Em um script do RouterOS, a sintaxe correta para atribuir valor a uma variável local e imprimi-la é:",
[":set $var \"olá\"; put $var","local var=\"olá\"; print var","var = 'olá'; echo var",":local var \"olá\"; :put $var",":declare var=olá; :print var"],3,
"No scripting RouterOS, ':local' declara variável local, atribuição por espaço e uso com $. ':put' imprime na saída.","mikrotik"))

Q.append(("O limite padrão máximo para o número de conexões rastreadas (connection tracking) no RouterOS depende:",
["Da licença (Level).","Da memória RAM do equipamento, sendo ajustado automaticamente.","Do número de usuários cadastrados.","Do protocolo utilizado (TCP/UDP).","Do plano de IP contratado."],1,
"O parâmetro 'total-max-entries' do conntrack é calculado em função da RAM disponível no dispositivo, podendo ser ajustado em /ip firewall connection tracking.","mikrotik"))

Q.append(("A função 'Address List' no firewall do RouterOS permite:",
["Armazenar rotas estáticas.","Agrupar endereços IP/redes em listas nomeadas para uso em regras (src/dst-address-list).","Listar endereços MAC bloqueados automaticamente.","Armazenar endereços MAC de clientes DHCP apenas.","Trocar endereços dinamicamente."],1,
"As address-list agrupam prefixos em nomes lógicos, usadas por 'src-address-list' e 'dst-address-list' em regras de firewall/mangle.","mikrotik"))

Q.append(("O protocolo EoIP do Mikrotik caracteriza-se por:",
["Encapsular Ethernet em TCP.","Encapsular Ethernet em IP usando GRE, criando um túnel L2 entre dois roteadores Mikrotik.","Funcionar apenas sobre IPv6.","Substituir o IPSec.","Fornecer apenas transporte de voz."],1,
"EoIP (Ethernet over IP) transporta frames Ethernet dentro de pacotes IP via GRE, viabilizando bridging entre sites distintos.","mikrotik"))

Q.append(("Para permitir acesso ao Winbox apenas a partir do IP 10.0.0.5, restringindo os demais, é correto configurar em '/ip service':",
["set winbox address=10.0.0.5/32","set winbox port=10.0.0.5","set winbox disabled=yes only=10.0.0.5","set winbox allowed=10.0.0.5","set winbox access=single 10.0.0.5"],0,
"Em '/ip service' pode-se limitar cada serviço por 'address'; somente IPs/prefixos listados têm acesso.","mikrotik"))

Q.append(("No RouterOS 7.x, o firewall IPv6 é configurado em:",
["/ip firewall filter","/ipv6 firewall filter","/firewall v6","/ip v6 fw","/routing v6"],1,
"Para IPv6 usa-se o ramo /ipv6 com chains análogas (input, output, forward).","mikrotik"))

Q.append(("O comando '/tool ping 8.8.8.8 count=5 interface=ether1' produz como efeito:",
["Envia 5 ICMP echo request para 8.8.8.8 utilizando ether1 como interface de saída.","Pinga o gateway da ether1.","Envia 5 pacotes TCP SYN.","Executa um traceroute.","Testa a ethernet com loopback."],0,
"/tool ping é o ping do RouterOS; 'count' limita o número de pacotes e 'interface' força a interface de origem.","mikrotik"))

Q.append(("Quanto à funcionalidade Bridge Filter no RouterOS, é correto afirmar que:",
["Filtra somente pacotes IPv4.","Permite aplicar regras em nível de camada 2 (frames) em interfaces bridge, análogas ao filter da L3.","Substitui o /ip firewall filter.","Age apenas em VLANs.","É aplicada somente em roteamento dinâmico."],1,
"O Bridge Filter opera em L2, filtrando frames que atravessam bridges (similar ao ebtables do Linux).","mikrotik"))

Q.append(("O recurso de 'PCQ' (Per Connection Queue) em Queue Types do RouterOS é utilizado para:",
["Dividir dinamicamente a banda entre múltiplos usuários/conexões de forma equitativa.","Priorizar apenas o protocolo ICMP.","Substituir o mangle.","Monitorar logs.","Descobrir vizinhos."],0,
"PCQ divide a banda dinamicamente por fluxo (ex: por src-address), garantindo equidade entre múltiplos usuários.","mikrotik"))

Q.append(("No Mikrotik, a ordem padrão das chains do firewall filter para pacotes em trânsito atravessando o roteador é:",
["input → output → forward","prerouting → input → forward → output → postrouting","Não se aplica, pois o filter só tem 'forward'.","Somente 'forward' é atravessado, precedido por prerouting e seguido por postrouting no /ip firewall mangle.","output → forward → input"],3,
"Pacotes em trânsito passam por prerouting (mangle/nat), decisão de rota, forward (mangle/filter) e postrouting (mangle/nat). No filter, tais pacotes batem apenas em 'forward'.","mikrotik"))

Q.append(("A função 'Netinstall' do Mikrotik é utilizada para:",
["Atualizar pacotes via internet.","Instalar/Reinstalar o RouterOS em um dispositivo via BOOTP/TFTP pela rede, apagando a configuração atual.","Configurar roteamento estático em massa.","Realizar backups automáticos na nuvem.","Configurar DHCP do usuário."],1,
"O Netinstall é utilitário Windows que reinstala o RouterOS via rede (PXE/BOOTP), sendo útil em recuperação e instalação inicial.","mikrotik"))

Q.append(("Ao realizar balanceamento de carga com dois links ISPs no Mikrotik, uma abordagem clássica é:",
["Configurar duas rotas estáticas com a mesma distância (ECMP), opcionalmente em conjunto com marcação mangle para roteamento por políticas.","Habilitar OSPF entre os dois ISPs.","Usar apenas NAT para alternar conexões.","Criar duas VLANs diferentes por ISP.","Configurar VRRP entre os provedores."],0,
"ECMP (rotas com mesmo dst e diferentes gateways) é uma forma simples. Para isolamento por origem (PCC/mark-routing) usa-se mangle + rotas marcadas.","mikrotik"))

Q.append(("Sobre o 'Safe Mode' acessado via Ctrl+X no terminal do RouterOS, é correto afirmar:",
["Ativa criptografia das senhas.","Permite testar alterações e desfazê-las automaticamente caso a sessão seja perdida antes da confirmação.","Desabilita o firewall.","Cria um backup automático.","Reinicia o roteador a cada comando."],1,
"O Safe Mode reverte automaticamente todas as alterações feitas durante a sessão, caso a conexão caia ou o modo seja abandonado de forma anormal.","mikrotik"))

Q.append(("Para criar um cliente DHCP no Mikrotik na interface ether1, utiliza-se:",
["/ip dhcp-client add interface=ether1 disabled=no","/ip address add interface=ether1 dhcp=yes","/ip dhcp-server add ether1","/ip client dhcp ether1","/ip pool add interface=ether1"],0,
"Adiciona-se a instância em /ip dhcp-client, associada à interface que receberá o endereço dinamicamente.","mikrotik"))

Q.append(("No RouterOS, as ferramentas '/tool torch' e '/tool traffic-monitor' são usadas, respectivamente, para:",
["Monitorar tráfego em tempo real de uma interface (por IP/porta/protocolo) e disparar ações quando o tráfego cruza limites definidos.","Realizar atualização de firmware e reiniciar o sistema.","Configurar VLAN e VPN.","Exportar configuração e importar certificados.","Testar disco e memória."],0,
"Torch mostra fluxos em tempo real; Traffic Monitor dispara scripts quando o tráfego atinge thresholds em uma interface.","mikrotik"))

# ----------------- LINUX CENTOS (33 questões) -----------------
Q.append(("No CentOS 7/8, o gerenciador de pacotes de alto nível padrão, que resolve dependências automaticamente, é:",
["dpkg","apt","yum (ou dnf, no CentOS 8)","pacman","zypper"],2,
"No CentOS 7 usa-se yum; no CentOS 8/Stream o dnf (com 'yum' mantido como link simbólico para compatibilidade). Ambos resolvem dependências via repositórios.","linux"))

Q.append(("O comando 'rpm -qa | grep httpd' tem como função:",
["Instalar o pacote httpd.","Listar todos os pacotes RPM instalados no sistema filtrando por 'httpd'.","Remover o pacote httpd.","Verificar assinaturas do pacote httpd.","Exibir apenas o arquivo de configuração do Apache."],1,
"'rpm -qa' (query all) lista todos os pacotes; o 'grep httpd' filtra pelos que contêm 'httpd' no nome.","linux"))

Q.append(("No systemd (CentOS 7+), para habilitar o serviço httpd na inicialização e iniciá-lo imediatamente, utiliza-se:",
["chkconfig httpd on; service httpd start","systemctl enable --now httpd","systemctl start httpd.daemon","service httpd enable","systemd --start=httpd"],1,
"O parâmetro '--now' do systemctl combina 'enable' (criar link em multi-user.target.wants) com 'start' em uma única chamada.","linux"))

Q.append(("A permissão representada por '755' em um arquivo significa:",
["rwxr-xr-x (dono: rwx; grupo: r-x; outros: r-x)","rw-rw-rw-","rwxrwxrwx","r--r--r--","rwxrw-r--"],0,
"7 = rwx; 5 = r-x; 5 = r-x. Portanto owner=rwx, group=r-x, others=r-x.","linux"))

Q.append(("No CentOS, o arquivo de configuração principal do Apache HTTP Server é:",
["/etc/apache/apache.conf","/etc/httpd/conf/httpd.conf","/etc/nginx/nginx.conf","/etc/apache2/apache2.conf","/var/www/conf/httpd.conf"],1,
"No pacote httpd do CentOS, o arquivo principal é /etc/httpd/conf/httpd.conf, com extensões em /etc/httpd/conf.d/.","linux"))

Q.append(("O comando para verificar se a porta 443 está sendo escutada localmente, utilizando o ss, é:",
["ss -tulpn | grep :443","ss -port 443","ss --listening=443","netstat -Z 443","ss -k 443"],0,
"'ss -tulpn' lista sockets TCP/UDP em listening, com PID/processo e sem resolução DNS. O grep filtra pela porta 443.","linux"))

Q.append(("No SELinux, o modo que registra violações de política no log mas não as bloqueia é:",
["enforcing","permissive","disabled","audit-only","targeted"],1,
"'Permissive' apenas registra (em /var/log/audit/audit.log) sem bloquear; 'enforcing' bloqueia; 'disabled' desativa totalmente. 'targeted' é uma política, não um modo."," linux"))

Q.append(("No firewalld, para liberar permanentemente o serviço http na zona pública, utiliza-se:",
["firewall-cmd --zone=public --add-service=http --permanent && firewall-cmd --reload","iptables -A INPUT -p tcp --dport 80 -j ACCEPT","systemctl allow http","firewall-cmd --open http --zone=public --save","firewalld --permit=http"],0,
"A liberação permanente exige '--permanent' seguido de '--reload' para que a regra entre em vigor no runtime.","linux"))

Q.append(("O arquivo de zona principal do BIND que define as configurações globais do servidor DNS é:",
["/etc/named.conf","/var/named/named.conf","/etc/bind/named.conf.local","/etc/dns/named.conf","/etc/resolv.conf"],0,
"No pacote bind do CentOS, a configuração principal fica em /etc/named.conf. Os arquivos de zona (ex.: db.example.com) ficam em /var/named/.","linux"))

Q.append(("Sobre o comando 'chmod g+s /srv/dados', é correto afirmar:",
["Atribui SUID ao diretório.","Atribui SGID ao diretório, fazendo com que novos arquivos criados herdem o grupo do diretório.","Torna o diretório imutável.","Remove permissões do grupo.","Ativa ACL no diretório."],1,
"O bit SGID em um diretório faz com que arquivos criados dentro dele sejam criados com o grupo do diretório (facilita compartilhamento).","linux"))

Q.append(("No shell Bash, o comando 'for i in $(seq 1 5); do echo $i; done' imprime:",
["Apenas o valor 5.","Os números de 1 a 5, um por linha.","Um erro de sintaxe.","Os números de 1 a 4.","Uma string literal '1 5'."],1,
"O for itera a sequência gerada por 'seq 1 5' (1,2,3,4,5), imprimindo cada número por linha.","linux"))

Q.append(("Para exibir as últimas 50 linhas do arquivo /var/log/messages e continuar acompanhando-o em tempo real, utiliza-se:",
["head -n 50 /var/log/messages","tail -n 50 -f /var/log/messages","cat /var/log/messages | more","less /var/log/messages","grep -f 50 /var/log/messages"],1,
"'tail -n 50' mostra as últimas 50 linhas e '-f' mantém o comando ativo acompanhando o crescimento do arquivo.","linux"))

Q.append(("Qual comando cria um novo usuário 'joao' com diretório /home/joao e shell /bin/bash?",
["adduser --home=/home/joao joao","useradd -m -d /home/joao -s /bin/bash joao","user add joao home=/home/joao","newuser joao /home/joao","mkuser joao bash"],1,
"'useradd -m -d <dir> -s <shell> <user>' cria o usuário com diretório home e shell especificados.","linux"))

Q.append(("O serviço NFS no CentOS é habilitado via pacote:",
["nfs-kernel-server","nfs-utils","nfsd","libnfs","nfs-common"],1,
"No CentOS/RHEL o pacote é 'nfs-utils', que fornece daemons e utilitários (exportfs, mount.nfs, etc.).","linux"))

Q.append(("A respeito do arquivo /etc/fstab, é correto afirmar que:",
["É um script executado na inicialização pelo cron.","Contém informações sobre os sistemas de arquivos a serem montados, usado por mount -a e na inicialização.","Configura o gerenciador de pacotes.","Define variáveis de ambiente globais.","Controla permissões SELinux."],1,
"O /etc/fstab define dispositivos/sistemas de arquivos que devem ser montados automaticamente na inicialização ou via 'mount -a'.","linux"))

Q.append(("Qual comando permite ver os processos em execução de forma contínua, com informações de CPU e memória?",
["ps aux","top","dmesg","iostat","uptime"],1,
"'top' atualiza dinamicamente processos, CPU, memória. 'ps aux' dá um snapshot.","linux"))

Q.append(("No CentOS 7+, o comando para verificar o status de um serviço chamado sshd é:",
["systemctl status sshd","service status sshd","systemd status sshd","ps sshd status","init sshd status"],0,
"No systemd usa-se 'systemctl status <unit>'. 'service' ainda é compatível, mas redireciona internamente.","linux"))

Q.append(("O arquivo /etc/passwd contém:",
["Apenas as senhas dos usuários em formato hash.","Informações dos usuários, incluindo UID, GID, diretório home e shell, sendo legível a todos os usuários.","Apenas os usuários com privilégios de root.","Os logs de login.","Lista de permissões sudo."],1,
"O /etc/passwd é público e contém uid:gid:gecos:home:shell. As senhas (hash) ficam em /etc/shadow (acesso restrito).","linux"))

Q.append(("Qual o diretório padrão, no CentOS, para publicação de conteúdo do servidor Apache?",
["/etc/httpd/","/var/www/html","/srv/http","/home/apache","/usr/share/apache"],1,
"O DocumentRoot default do Apache no CentOS/RHEL é /var/www/html.","linux"))

Q.append(("Para verificar se o SELinux está em modo enforcing ou permissive, utiliza-se:",
["selinux status","getenforce","semodule -s","sestatus -e","checkpolicy"],1,
"'getenforce' exibe Enforcing/Permissive/Disabled. 'sestatus' (sem -e) dá visão detalhada; 'setenforce 0/1' alterna os modos em tempo de execução.","linux"))

Q.append(("No CentOS, o Samba é utilizado principalmente para:",
["Servir arquivos via NFS.","Compartilhamento de arquivos e impressoras compatível com o protocolo SMB/CIFS utilizado por clientes Windows.","Servidor SSH.","Servidor FTP seguro.","Gerenciador de e-mails."],1,
"Samba implementa SMB/CIFS, permitindo que Linux ofereça compartilhamentos acessíveis por clientes Windows/macOS/Linux.","linux"))

Q.append(("A sintaxe correta de uma linha no cron para executar /usr/local/bin/backup.sh todos os dias às 02:30 é:",
["30 2 * * * /usr/local/bin/backup.sh","2 30 * * * /usr/local/bin/backup.sh","* * 2 30 * /usr/local/bin/backup.sh","@daily 02:30 /usr/local/bin/backup.sh","30 2 */1 * * /usr/local/bin/backup.sh"],0,
"Ordem do cron: minuto hora dia-mês mês dia-semana. Para 02:30 diariamente: '30 2 * * *'.","linux"))

Q.append(("Qual comando monta um compartilhamento NFS remoto 'servidor:/dados' em /mnt/dados?",
["nfs-mount servidor:/dados /mnt/dados","mount -t nfs servidor:/dados /mnt/dados","mount /dev/servidor /mnt/dados","mount -o cifs servidor:/dados /mnt/dados","nfs connect servidor /mnt/dados"],1,
"A sintaxe do mount é 'mount -t nfs <host>:<path> <mountpoint>'. O tipo 'cifs' é usado para SMB.","linux"))

Q.append(("No iptables, a cadeia utilizada para filtragem de pacotes destinados ao próprio host é:",
["FORWARD","OUTPUT","INPUT","PREROUTING","POSTROUTING"],2,
"Na tabela filter: INPUT (destino local), OUTPUT (origem local), FORWARD (em trânsito). PREROUTING/POSTROUTING são de outras tabelas (nat, mangle).","linux"))

Q.append(("Sobre o comando 'journalctl -u sshd -xe', é correto afirmar que:",
["Apaga os logs do sshd.","Exibe os logs do serviço sshd com contexto (-x) e posiciona-se no fim (-e).","Reinicia o sshd.","Exporta os logs do sshd para /var/log/sshd.","Habilita o sshd no boot."],1,
"O journalctl consulta o journal do systemd. '-u' filtra por unidade, '-x' inclui explicações, '-e' pula para o fim.","linux"))

Q.append(("Qual diretiva do arquivo /etc/ssh/sshd_config, quando definida como 'no', impede o login direto como usuário root via SSH?",
["RootLogin no","PermitRootLogin no","DenyRoot yes","AllowRootLogin no","DisableRoot yes"],1,
"A diretiva correta é 'PermitRootLogin no'. Para permitir apenas com chave e sem senha usa-se 'prohibit-password'.","linux"))

Q.append(("Qual comando aplica configurações após editar o /etc/sysctl.conf sem reiniciar?",
["sysctl -w","sysctl -p","sysctl reload","sysctl --refresh","sysctl restart"],1,
"'sysctl -p' carrega os parâmetros do arquivo (por padrão /etc/sysctl.conf ou arquivos em /etc/sysctl.d).","linux"))

Q.append(("No CentOS 8/Stream, qual comando exibe repositórios habilitados?",
["yum list repos","dnf repolist","rpm --repos","dnf show repos","yum --list-enabled"],1,
"'dnf repolist' (ou 'dnf repolist --all') mostra os repositórios. Em CentOS 7, 'yum repolist' tem comportamento equivalente.","linux"))

Q.append(("No shell script Bash, a estrutura 'if [ -f \"$arq\" ]; then ... fi' avalia:",
["Se $arq é um diretório.","Se $arq é um link simbólico.","Se $arq é um arquivo regular existente.","Se $arq está vazio.","Se $arq é um dispositivo."],2,
"O teste '-f' retorna verdadeiro se o caminho existir e for um arquivo regular. '-d' testa diretório, '-L' link, '-s' tamanho >0.","linux"))

Q.append(("Para alterar o dono de /var/www/html/site e seu conteúdo recursivamente para o usuário apache e grupo apache, usa-se:",
["chmod -R apache:apache /var/www/html/site","chown apache.apache /var/www/html/site","chown -R apache:apache /var/www/html/site","chgrp apache -R /var/www/html/site","setfacl -R apache /var/www/html/site"],2,
"'chown -R usuario:grupo caminho' altera dono/grupo recursivamente. 'chmod' altera permissões, não proprietário.","linux"))

Q.append(("O comando que permite visualizar de forma contínua o uso de I/O por processo no Linux é:",
["top","iotop","vmstat","sar -b","ps -io"],1,
"'iotop' exibe I/O por processo em tempo real, similar ao 'top'. Requer privilégios de root.","linux"))

Q.append(("No servidor DHCP do CentOS (dhcpd), a diretiva que define o pool de endereços para um escopo é:",
["pool-ip","range","dhcp-range","address-pool","scope"],1,
"Dentro de um bloco subnet do dhcpd.conf usa-se 'range <ip_inicial> <ip_final>;' para definir a faixa distribuída.","linux"))

Q.append(("Sobre LVM (Logical Volume Manager), é correto afirmar:",
["Substitui o sistema de arquivos ext4.","Permite criar volumes lógicos redimensionáveis sobre um ou mais volumes físicos (PV) agrupados em volume groups (VG).","É exclusivo para discos SSD.","É incompatível com XFS.","Funciona apenas como RAID 0."],1,
"No LVM: Physical Volume (PV) → Volume Group (VG) → Logical Volume (LV). LVs podem ser expandidos/reduzidos e usados com qualquer FS suportado (ext4, xfs, etc.).","linux"))

# Garantia de 100 questões
assert len(Q) == 100, f"Esperado 100, obtido {len(Q)}"

# Embaralha preservando reprodutibilidade
random.seed(42)
random.shuffle(Q)

# =======================================================================
# GERAÇÃO DO DOCX
# =======================================================================
doc = Document()

# Margens
sec = doc.sections[0]
sec.top_margin = Cm(2.0)
sec.bottom_margin = Cm(2.0)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.0)

# Estilo base
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)

# Cabeçalho
header = sec.header.paragraphs[0]
header.text = "Concurso Público – Analista de Informática"
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
for r in header.runs:
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

# Rodapé com número de página
footer = sec.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run()
fld1 = run._r.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'}); run._r.append(fld1)
run2 = footer.add_run()
it = run2._r.makeelement(qn('w:instrText'), {}); it.text = ' PAGE '; run2._r.append(it)
run3 = footer.add_run()
fld2 = run3._r.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'}); run3._r.append(fld2)

# ----- Título principal -----
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("PROVA – ANALISTA DE INFORMÁTICA")
r.bold = True; r.font.size = Pt(18); r.font.color.rgb = RGBColor(0x1a, 0x36, 0x5d)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Cisco IOS, Mikrotik/RouterOS e GNU/Linux (CentOS)")
r.italic = True; r.font.size = Pt(13); r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Banca Examinadora: estilo Vunesp")
r.font.size = Pt(11); r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph("")

# ----- Instruções -----
pi = doc.add_paragraph()
ri = pi.add_run("INSTRUÇÕES GERAIS"); ri.bold = True; ri.font.size = Pt(13)

instr = [
    "1. A prova é composta de 100 (cem) questões objetivas, cada uma com 5 (cinco) alternativas (A, B, C, D, E), das quais apenas uma é correta.",
    "2. A duração total da prova é de 4 (quatro) horas, já incluído o tempo para preenchimento da folha de respostas.",
    "3. Não é permitido o uso de calculadora, celular, materiais de consulta ou qualquer meio eletrônico durante a prova.",
    "4. Leia atentamente cada questão antes de assinalar a alternativa escolhida; rasuras invalidam a marcação.",
    "5. As questões abordam, em distribuição aproximada, os seguintes temas: Cisco IOS, Mikrotik/RouterOS e GNU/Linux (CentOS).",
    "6. O gabarito comentado encontra-se ao final deste caderno.",
]
for t in instr:
    pr = doc.add_paragraph(t); pr.paragraph_format.space_after = Pt(2)

doc.add_paragraph("")
sep = doc.add_paragraph()
sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
rs = sep.add_run("— QUESTÕES —"); rs.bold = True; rs.font.size = Pt(12)
doc.add_paragraph("")

# ----- Questões -----
LETRAS = ['A','B','C','D','E']
for idx, (enun, alts, _cor, _com, _tema) in enumerate(Q, start=1):
    # Enunciado
    pq = doc.add_paragraph()
    pq.paragraph_format.space_before = Pt(6)
    pq.paragraph_format.space_after = Pt(3)
    rn = pq.add_run(f"Questão {idx:03d}. "); rn.bold = True
    # Suporta quebras de linha literais "\n" no enunciado
    for i, ln in enumerate(enun.split("\n")):
        if i > 0:
            pq.add_run().add_break()
        pq.add_run(ln)
    # Alternativas
    for li, alt in enumerate(alts):
        pa = doc.add_paragraph()
        pa.paragraph_format.left_indent = Cm(0.8)
        pa.paragraph_format.space_after = Pt(1)
        ra = pa.add_run(f"({LETRAS[li]}) "); ra.bold = True
        pa.add_run(alt)

# ----- Quebra de página -----
doc.add_page_break()

# ----- Gabarito comentado -----
pg = doc.add_paragraph()
pg.alignment = WD_ALIGN_PARAGRAPH.CENTER
rg = pg.add_run("GABARITO COMENTADO")
rg.bold = True; rg.font.size = Pt(16); rg.font.color.rgb = RGBColor(0x1a, 0x36, 0x5d)
doc.add_paragraph("")

# Tabela-resumo do gabarito (5 colunas x 20 linhas)
tbl = doc.add_table(rows=20, cols=10)
tbl.style = 'Light Grid Accent 1'
for i in range(100):
    r = i % 20
    c = (i // 20) * 2
    cell_n = tbl.cell(r, c)
    cell_a = tbl.cell(r, c+1)
    cell_n.text = f"{i+1:03d}"
    cell_a.text = LETRAS[Q[i][2]]
    for cc in (cell_n, cell_a):
        for pp in cc.paragraphs:
            for rr in pp.runs:
                rr.font.size = Pt(9)
                if cc is cell_a:
                    rr.bold = True

doc.add_paragraph("")

# Comentários detalhados
pt = doc.add_paragraph()
rt = pt.add_run("Comentários e Justificativas")
rt.bold = True; rt.font.size = Pt(13); rt.font.color.rgb = RGBColor(0x1a, 0x36, 0x5d)

for idx, (enun, alts, cor, com, tema) in enumerate(Q, start=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"Questão {idx:03d} – Gabarito: {LETRAS[cor]}")
    r.bold = True
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Cm(0.5)
    p2.paragraph_format.space_after = Pt(4)
    r2 = p2.add_run(f"Alternativa correta: ({LETRAS[cor]}) {alts[cor]}")
    r2.italic = True
    p3 = doc.add_paragraph()
    p3.paragraph_format.left_indent = Cm(0.5)
    p3.paragraph_format.space_after = Pt(8)
    p3.add_run("Comentário: ").bold = True
    p3.add_run(com)

out = "/home/ubuntu/prova_cisco_mikrotik_linux_vunesp.docx"
doc.save(out)
print("OK:", out, "- questões:", len(Q))
